import base64
import io
import json
import os
import re
import tempfile
from datetime import datetime

import streamlit as st
from openai import OpenAI
from PIL import Image, ImageStat

try:
    from docx import Document
    from docx.shared import Inches
except Exception:
    Document = None
    Inches = None


APP_VERSION = "V2.0.1"
MEMORY_FILE = "geo_memory.jsonl"


# ================= 全局初始化：创建轻量记忆库，方便后续做错误案例和专家修正沉淀 =================
if not os.path.exists(MEMORY_FILE):
    with open(MEMORY_FILE, "w", encoding="utf-8") as f:
        f.write("")


# ================= Secrets 读取 =================
def get_api_key():
    """优先从 Streamlit secrets 读取，其次才使用侧边栏输入。"""
    try:
        if "OPENAI_API_KEY" in st.secrets:
            return st.secrets["OPENAI_API_KEY"]
    except Exception:
        pass
    return ""


# ================= 页面配置 =================
st.set_page_config(
    page_title="地质AI辅助解译系统",
    page_icon="🪨",
    layout="wide",
)


# ================= 记忆库与工具函数 =================
def load_memory_examples(limit=3):
    """读取最近几条记忆，用于提示词增强。"""
    examples = []
    if not os.path.exists(MEMORY_FILE):
        return examples
    with open(MEMORY_FILE, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            try:
                examples.append(json.loads(line))
            except Exception:
                continue
    return examples[-limit:]



def save_memory_record(record):
    """把本次解译结果与人工修正写入轻量记忆库。"""
    with open(MEMORY_FILE, "a", encoding="utf-8") as f:
        f.write(json.dumps(record, ensure_ascii=False) + "\n")



def get_image_quality(image_bytes):
    """基础图片质量控制，帮助系统先判断是否值得精细解译。"""
    img = Image.open(io.BytesIO(image_bytes)).convert("RGB")
    width, height = img.size
    stat = ImageStat.Stat(img)
    brightness = sum(stat.mean) / 3
    sharpness_proxy = sum(stat.var) / 3
    return {
        "图像宽度": width,
        "图像高度": height,
        "平均亮度": round(brightness, 2),
        "清晰度": round(sharpness_proxy, 2),
        "图像太小": width < 800 or height < 600,
        "图像太暗": brightness < 35,
        "图像太亮": brightness > 220,
        "图像太模糊": sharpness_proxy < 250,
    }



def compress_image_for_inference(image_bytes, max_side=1600, quality=85):
    """压缩并轻微缩放图片，减少上传体积并尽量保留地质细节。"""
    img = Image.open(io.BytesIO(image_bytes)).convert("RGB")
    original_width, original_height = img.size

    if max(original_width, original_height) > max_side:
        img.thumbnail((max_side, max_side), Image.Resampling.LANCZOS)

    output = io.BytesIO()
    img.save(output, format="JPEG", quality=quality, optimize=True, progressive=True)
    return output.getvalue(), img.size



def classify_image_type_prompt():
    """返回统一的图像分类与解译要求。"""#规则说明书
    return """
你是地质技能大赛野外AI赛道助手。请先判断图像类型,再输出结构化结果。

图像类型可取值：
- 无人机航拍宏观地质影像
- 地面近景地质影像
- 宏微观混合地质影像
- 地质剖面图
- 图像类型不明

要求：
1. 先输出一段自然语言，且第一句必须是“这是一张……图。”
2. 再输出严格 JSON,方便程序解析。
3. JSON 中必须包含以下字段（每一项要换行输出）：
   - 图像类型
   - 图像质量评估
   - 可见要素
   - 推荐任务
   - 不适合任务
   - 自然语言解译
   - 结构化解译
   - 置信度
   - 不确定性
   - 下一步建议
4. 如果是 无人机航拍宏观地质图，请重点关注地形、地貌单元、地层展布、线性构造、沟谷水系，岩性分布
5. 如果是 地面实拍微观地质图，请重点关注岩性、结构构造、粒度、层理、节理裂隙、风化。
6. 如果是 宏微观混合地质影像，请分别给出宏观和微观两个分析块。
7. 如果是 地质剖面图，请重点关注地层和岩层特征、地质构造、岩浆活动、岩性、地质历史。
8. 如果图像质量差，请明确指出限制，不要强行下结论。
9. 语言要专业、克制、符合比赛规则，不要编造不可见信息。
10. 不要输出 Markdown。
""".strip()



def build_prompt(memory_examples, user_focus):
    """把记忆库和本次任务合成提示词。"""
    memory_text = "暂无历史记忆。"
    if memory_examples:
        memory_text = "\n".join(
            [
                f"- 案例{idx + 1}：输入特征={item.get('summary', '无')}；最终结论={item.get('final_conclusion', '无')}"
                for idx, item in enumerate(memory_examples)
            ]
        )

    return f"""
{classify_image_type_prompt()}

历史记忆参考：
{memory_text}

本次用户关注点：{user_focus}

补充要求：
- 自然语言解译必须是适合直接展示给用户的中文自然语言。
- 结构化解译必须是 JSON 对象，按图像类型动态组织字段。
- 如果是 无人机航拍宏观地质影像，结构化解译中建议包含：地形地貌、地层展布、线性构造、沟谷水系、岩性分布。
- 如果是 地面近景地质影像，结构化解译中建议包含：岩性、结构构造、粒度、层理、节理裂隙、风化。
- 如果是 地质剖面图，结构化解译中建议包含：地层接触关系、产状、褶皱、断层、岩体接触、时代关系。
- 如果是 宏微观混合地质影像，结构化解译中建议包含：宏观和微观两个分析块。
- 置信度用 0 到 1 之间的小数表示。
""".strip()



def extract_json_from_text(text):
    """尽量从模型输出中提取 JSON,容错处理多余文本。"""
    try:
        return json.loads(text)
    except Exception:
        match = re.search(r"\{.*\}", text, re.S)
        if match:
            try:
                return json.loads(match.group(0))
            except Exception:
                return None
    return None



def make_mock_result(image_quality):
    """离线模式下返回标准化结果，便于比赛现场演示。"""
    return {
        "图像类型": "地面近景地质影像",
        "图像质量评估": image_quality,
        "可见要素": ["岩面", "节理裂隙", "轻微风化面"],
        "推荐研究任务": ["岩性识别", "风化特征识别", "构造特征识别"],
        "不适合研究任务": ["精确层位追踪", "高精度产状测量"],
        "自然语言解译": "这是一张地面实拍微观地质图。图像更适合观察岩性、风化面、节理裂隙和局部构造特征，但由于缺少尺度参照，细粒度判断仍需结合补拍照片和野外背景信息。",
        "结构化解译": {
            "岩性": "中细粒花岗岩（模拟示例）",
            "结构构造": "中细粒花岗结构，块状构造",
            "矿物或成分": ["石英", "钾长石", "斜长石", "少量黑云母"],
            "风化程度": "表面轻微风化，节理裂隙较发育",
            "地质解译": "整体表现为酸性侵入岩特征，适合做岩性与风化识别示范。",
        },
        "置信度": 0.72,
        "不确定性": ["缺少比例尺", "缺少拍摄位置和区域地质背景"],
        "下一步建议": ["补充近景与远景各一张", "增加比例尺或地质锤参照", "记录拍摄位置与层位信息"],
        "模型原始输出": "",
    }


# ================= 模型调用：保留流式输出，同时在底层解析结构化结果 =================
def analyze_geology_image(image_path, user_prompt, api_key, model_name):
    """调用大模型(通义千问VL),返回流式文本生成器。"""
    if not api_key:#检查API key
        yield "⚠️ 请先在左侧边栏输入有效的 API Key。"
        return

    client = OpenAI(
        api_key=api_key,
        base_url="https://dashscope.aliyuncs.com/compatible-mode/v1",
    )#初始化大模型(可以配置不同大模型)

    try:#读取图片并转为Base64
        with open(image_path, "rb") as image_file:
            base64_image = base64.b64encode(image_file.read()).decode("utf-8")
    except FileNotFoundError:
        yield "❌ 找不到图片文件，请检查上传。"
        return

    messages = [#构造多模态请求消息
        {
            "role": "user",
            "content": [
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"},
                },
                {
                    "type": "text",
                    "text": user_prompt,
                },
            ],
        }
    ]

    try:#发起流式请求
        completion = client.chat.completions.create(
            model=model_name,#选择大模型
            messages=messages,
            temperature=0.2,#地质解译需要严谨，降低随机性
            max_tokens=1800,#最长输出
            stream=True,#开启流式返回
        )
        for chunk in completion:#逐块Yield输出
            if chunk.choices and chunk.choices[0].delta.content:
                yield chunk.choices[0].delta.content
    except Exception as e:#异常处理
        yield f"❌ API调用异常: {str(e)}"



def render_result(result):#把结构化结果渲染成适合比赛展示的页面。
    """把结构化结果渲染成适合比赛展示的页面。"""
    st.subheader("3. 结构化解译结果")
    if not result:#检查结果
        st.info("请先完成一次解译。")
        return

    if result.get("error"):#检查错误
        st.error(result["error"])
        if result.get("raw_text"):#检查原始输出
            st.text_area("模型原始输出", result["raw_text"], height=180)
        return

    col_a, col_b, col_c = st.columns(3)#创建三个指标展示置信度、图像类型、推荐任务数
    col_a.metric("置信度", f"{round(result.get('置信度', 0) * 100)}%")
    col_b.metric("图像类型", result.get("图像类型", "图像类型不明"))
    col_c.metric("任务数", len(result.get("推荐任务", [])))

    st.text_area("自然语言解译\n" + result.get("自然语言解译", ""), height=120)
    st.write("可见要素")
    st.write(result.get("可见要素", []))
    st.write("推荐任务")
    st.write(result.get("推荐任务", []))
    st.write("不适合任务")
    st.write(result.get("不适合任务", []))
    st.write("不确定性")
    st.write(result.get("不确定性", []))
    st.write("下一步建议")
    st.write(result.get("下一步建议", []))

    with st.expander("查看底层结构化结果 JSON"):#查看底层结构化结果JSON（给专业人士或调试人员看的）
        st.json(result.get("结构化解译", {}))



def export_docx(image_bytes, result, output_path):#导出Word报告
    """导出 Word 报告；如果缺少依赖，则返回友好提示。"""
    if Document is None:#检查依赖
        return False, "当前环境缺少 python-docx，无法生成 .docx。"

    doc = Document()#创建Word文档
    doc.add_heading("地质AI辅助解译报告", level=1)#添加标题
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")#添加生成时间
    doc.add_paragraph(f"系统版本：{APP_VERSION}")#添加系统版本

    doc.add_heading("一、图像信息", level=2)
    doc.add_paragraph(f"图像类型：{result.get('图像类型', '图像类型不明')}")
    iq = result.get("图像质量评估", {})#获取图像质量评估
    if not isinstance(iq, dict):
        iq = {}#如果图像质量评估不是字典，则设置为空字典
    width = iq.get("图像宽度", iq.get("width", "NA"))#获取图像宽度
    height = iq.get("图像高度", iq.get("height", "NA"))#获取图像高度
    doc.add_paragraph(f"图像宽高：{width} × {height}")

    if image_bytes and Inches is not None:#检查图片
        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
            tmp.write(image_bytes)
            tmp_img_path = tmp.name
        try:#插入图片
            doc.add_picture(tmp_img_path, width=Inches(5.8))
        finally:#清理临时文件
            if os.path.exists(tmp_img_path):#删除临时文件
                os.unlink(tmp_img_path)

    doc.add_heading("二、自然语言结论", level=2)#添加自然语言结论
    doc.add_paragraph(result.get("自然语言解译", ""))

    doc.add_heading("三、结构化解译结果", level=2)#添加结构化解译结果
    structured = result.get("结构化解译", {})#获取结构化解译结果
    if isinstance(structured, dict):#如果结构化解译结果是字典
        for key, value in structured.items():#遍历结构化解译结果
            if isinstance(value, list):
                value_text = ", ".join(map(str, value))#如果结构化解译结果是列表，则将列表转换为字符串
            elif isinstance(value, dict):
                value_text = json.dumps(value, ensure_ascii=False)#如果结构化解译结果是字典，则将字典转换为字符串
            else:
                value_text = str(value)#如果结构化解译结果是其他类型，则将结果转换为字符串
            doc.add_paragraph(f"{key}：{value_text}")
    else:
        doc.add_paragraph(str(structured))#如果结构化解译结果不是字典，则将结果转换为字符串

    doc.add_heading("四、置信度与不确定性", level=2)
    doc.add_paragraph(f"置信度：{result.get('置信度', 0)}")#添加置信度

    uncertainty = result.get("不确定性", [])#获取不确定性
    if isinstance(uncertainty, list):
        uncertainty_text = ", ".join(map(str, uncertainty))#如果不确定性是列表，则将列表转换为字符串
    elif uncertainty is None:
        uncertainty_text = ""#如果不确定性是None，则设置为空字符串
    else:
        uncertainty_text = str(uncertainty)#如果不确定性是其他类型，则将结果转换为字符串    

    suggestions = result.get("下一步建议", [])
    if isinstance(suggestions, list):
        suggestions_text = ", ".join(map(str, suggestions))
    elif suggestions is None:
        suggestions_text = ""
    else:
        suggestions_text = str(suggestions)

    doc.add_paragraph(f"不确定性：{uncertainty_text}")
    doc.add_paragraph(f"建议：{suggestions_text}")

    doc.save(output_path)#保存文档
    return True, output_path#返回成功和路径 (Word报告路径)


# ================= 主界面 =================
st.title("🪨 地质AI辅助解译与成图系统")
st.caption("适合比赛演示的可用原型：保留流式输出、结构化结果、人工修正和 Word 导出。")
st.markdown("---")

# 侧边栏设置
st.sidebar.title("⚙️ 系统设置")
api_key = st.sidebar.text_input("输入API Key", type="password", help="请在阿里云百炼平台获取")
model_choice = st.sidebar.selectbox("选择AI大模型", ["qwen-vl-max", "qwen-vl-plus"])
use_mock_data = st.sidebar.checkbox("🚨 启用离线模拟模式 (断网备用)")
user_focus = st.sidebar.text_area("本次关注点", "请在此输入您本次的解译关注点")
st.sidebar.info(f"当前系统版本：{APP_VERSION}\n支持流式输出 + 结构化解析")

col1, col2 = st.columns([1, 1.25])

with col1:
    st.subheader("📥 1. 数据输入")
    uploaded_file = st.file_uploader("请上传野外露头或无人机照片", type=["jpg", "png", "jpeg"])
    image_bytes = None
    image_quality = None

    if uploaded_file is not None:#检查图片
        image_bytes = uploaded_file.getvalue()#获取图片字节
        compressed_image_bytes, compressed_size = compress_image_for_inference(image_bytes)#压缩图片用于推理
        st.image(image_bytes, caption="已上传地质图像", use_container_width=True)
        if compressed_size != Image.open(io.BytesIO(image_bytes)).size:
            st.caption(f"推理时已自动压缩为 {compressed_size[0]} × {compressed_size[1]}，以提升速度并减少传输体积。")
        image_quality = get_image_quality(image_bytes)#获取图片质量
        st.json(image_quality)#显示图片质量
        if image_quality["图像太小"] or image_quality["图像太暗"] or image_quality["图像太亮"] or image_quality["图像太模糊"]:
            st.warning("当前图片可能不适合精细解译，建议先补充更清晰的近景或增加尺度参照。")#提示用户图片质量不佳

    st.subheader("✏️ 4. 人工修正")
    manual_lithology = st.text_input("修正岩性", value=st.session_state.get("manual_lithology", ""))#修正岩性
    manual_comment = st.text_area("修正说明", value=st.session_state.get("manual_comment", ""), height=90)#修正说明
    if st.button("保存人工修正", use_container_width=True):#保存人工修正
        st.session_state.manual_lithology = manual_lithology
        st.session_state.manual_comment = manual_comment
        st.success("已保存人工修正，本次修正将进入记忆库。")

with col2:
    st.subheader("🤖 2. AI 智能解译")

    if "raw_stream_text" not in st.session_state:#初始化session state(会话状态)
        st.session_state.raw_stream_text = ""
    if "structured_result" not in st.session_state:
        st.session_state.structured_result = {}
    if "tmp_path" not in st.session_state:
        st.session_state.tmp_path = None

    if uploaded_file is not None:#如果已经上传了图片
        memory_examples = load_memory_examples(limit=3)#加载记忆库
        geo_prompt = build_prompt(memory_examples, user_focus)#构建提示词(由两个可变提示词(记忆库和用户关注点)和一个固定文本(规则说明书 )组成)

        if st.button("🚀 开始 AI 智能解译", type="primary", use_container_width=True):#开始解译按钮，点击后会清空旧结果
            st.session_state.raw_stream_text = ""
            st.session_state.structured_result = {}

            if use_mock_data:#如果启用离线模拟模式
                st.info("当前为离线模拟模式，未调用真实 API。")
                mock = make_mock_result(image_quality or {})
                st.session_state.raw_stream_text = mock.get("自然语言解译", "") + "\n" + json.dumps(mock, ensure_ascii=False)
                st.session_state.structured_result = mock
            else:
                api_key = get_api_key() or api_key
                if not api_key:#检查API key
                    st.error("请先在 Secrets 中配置 OPENAI_API_KEY，或在侧边栏输入 API Key。")
                    st.stop()

                with st.spinner("AI 正在进行结构化地质解译，请稍候..."):
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as tmp_file:
                        tmp_file.write(compressed_image_bytes)
                        st.session_state.tmp_path = tmp_file.name

                    try:
                        stream_gen = analyze_geology_image(
                            image_path=st.session_state.tmp_path,
                            user_prompt=geo_prompt,
                            api_key=api_key,
                            model_name=model_choice,
                        )#analyze_geology_image函数调用大模型进行解译
                        full_text = st.write_stream(stream_gen)#展示打字机效果并保存完整结果
                        st.session_state.raw_stream_text = full_text or ""#保存完整结果
                        parsed = extract_json_from_text(st.session_state.raw_stream_text)#extract_json_from_text函数提取JSON
                        if parsed is None:#检查JSON
                            st.session_state.structured_result = {
                                "error": "模型输出中未解析到有效 JSON",#错误信息
                                "模型原始输出": st.session_state.raw_stream_text,#原始  输出
                            }
                        else:
                            st.session_state.structured_result = parsed#保存结构化结果
                    finally:
                        if st.session_state.tmp_path and os.path.exists(st.session_state.tmp_path):#删除临时文件
                            os.unlink(st.session_state.tmp_path)#删除临时文件
                            st.session_state.tmp_path = None#清空临时文件路径

            # 把人工修正和本次结果写入记忆库，方便后续增强提示词
            if st.session_state.structured_result and not st.session_state.structured_result.get("error"):
                st.session_state.structured_result["manual_lithology"] = st.session_state.get("manual_lithology", "")#把人工修正写入结果对象
                st.session_state.structured_result["manual_comment"] = st.session_state.get("manual_comment", "")
                st.session_state.structured_result["image_quality_assessment"] = image_quality or {}#把图片质量写入结果对象

                save_memory_record(
                    {
                        "timestamp": datetime.now().isoformat(),#添加时间戳
                        "summary": (st.session_state.structured_result.get("自然语言解译") or "")[:120],#添加自然语言解译
                        "final_conclusion": json.dumps(st.session_state.structured_result.get("结构化解译", {}), ensure_ascii=False),#添加结构化解译
                        "manual_lithology": st.session_state.get("manual_lithology", ""),#添加人工修正
                        "manual_comment": st.session_state.get("manual_comment", ""),
                    }#保存记忆库
                )
                st.success("解译完成，结果已写入轻量记忆库。")#提示用户解译完成

        if st.session_state.raw_stream_text:#如果存在流式输出
            st.subheader("3. 流式输出原文")
            st.text_area("实时生成内容", st.session_state.raw_stream_text, height=180)

        if st.session_state.structured_result:
            render_result(st.session_state.structured_result)

# ================= 导出 Word 报告 =================
st.markdown("---")
if uploaded_file is not None and st.session_state.structured_result and not st.session_state.structured_result.get("error"):
    st.subheader("5. 报告导出")
    if st.button("📥 导出标准化地质报告 (Word)", use_container_width=True):
        with st.spinner("正在生成 Word 报告..."):
            tmp_docx = os.path.join(
                tempfile.gettempdir(),
                f"geology_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            )
            ok, result = export_docx(image_bytes, st.session_state.structured_result, tmp_docx)
            if ok:
                with open(result, "rb") as f:
                    st.download_button(
                        "点击下载 Word 报告",
                        data=f.read(),
                        file_name=os.path.basename(result),
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )
                st.toast("Word 报告已生成，可以下载。", icon="📄")
            else:
                st.warning(result)
