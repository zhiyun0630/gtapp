
import base64
import io
import json
import os
import re
import tempfile
import urllib.error
import urllib.request
from datetime import datetime
from pathlib import Path

import streamlit as st
from PIL import Image, ImageStat

from geology_database import GeologyDatabase, GeologyDatabaseError
from profile_generator import (
    ProfileGenerationError,
    draw_geological_profile,
    generate_profile_parameters,
    validate_profile_data,
)
from project_manager import ProjectManager, ProjectManagerError
from regional_analysis import RegionalAnalysisError, analyze_region, extract_json_object

try:
    from docx import Document
    from docx.shared import Inches
except Exception:
    Document = None
    Inches = None


APP_VERSION = "V4.0.0"
BASE_DIR = Path(__file__).resolve().parent
MEMORY_FILE = str(BASE_DIR / "geo_memory.jsonl")
PROJECTS_DIR = BASE_DIR / "projects"
PROJECT_MANAGER = ProjectManager(PROJECTS_DIR)


# ================= 全局初始化：创建轻量记忆库，方便后续做错误案例和专家修正沉淀 =================
if not os.path.exists(MEMORY_FILE):
    with open(MEMORY_FILE, "w", encoding="utf-8") as f:
        f.write("")


# ================= 页面配置 =================
st.set_page_config(
    page_title="地质AI辅助解译系统",
    page_icon="🪨",
    layout="wide",
)


# ================= 记忆库与工具函数 =================
def load_memory_examples(limit=5, memory_file=MEMORY_FILE):
    """读取最近几条记忆，用于提示词增强。"""
    examples = []
    if not os.path.exists(memory_file):
        return examples
    try:
        with open(memory_file, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                try:
                    examples.append(json.loads(line))
                except (json.JSONDecodeError, TypeError):
                    continue
    except OSError:
        return []
    return examples[-limit:]



def save_memory_record(record, memory_file=MEMORY_FILE):
    """把本次解译结果与人工修正写入轻量记忆库。"""
    Path(memory_file).parent.mkdir(parents=True, exist_ok=True)
    with open(memory_file, "a", encoding="utf-8") as f:
        f.write(json.dumps(record, ensure_ascii=False) + "\n")



def get_image_quality(image_bytes):
    """基础图片质量控制，帮助系统先判断是否值得精细解译。"""
    img = Image.open(io.BytesIO(image_bytes)).convert("RGB")
    width, height = img.size
    stat = ImageStat.Stat(img)
    brightness = sum(stat.mean) / 3
    sharpness_proxy = sum(stat.var) / 3
    return {
        "图像宽度": width,
        "图像高度": height,
        "平均亮度": round(brightness, 2),
        "清晰度": round(sharpness_proxy, 2),
        "图像太小": width < 800 or height < 600,
        "图像太暗": brightness < 35,
        "图像太亮": brightness > 220,
        "图像太模糊": sharpness_proxy < 250,
    }



def classify_image_type_prompt():
    """返回六类竞赛图像的统一分类和地质安全规则。"""
    return """
你是地质技能竞赛的“野外地质观察助手”。先判断图像类型，再提取可见事实，只能给出有限度的初步解释。

图像类型可取值：
- 无人机航拍地貌影像
- 地面露头照片
- 多视角露头照片组合
- 地质剖面图
- 三维模型截图
- 地质图

分类观察重点：
1. 无人机航拍地貌影像：地貌单元、山脊、沟谷、坡面、地质体展布和构造线索。
2. 地面露头照片：颜色、结构、构造、粒度、风化和节理。
3. 多视角露头照片组合：多照片一致性、空间连续性、延伸方向和露头规模。
4. 地质剖面图：地层、接触关系、断层、褶皱和岩浆活动。
5. 三维模型截图：地形关系、空间展布和地貌关系；不得声称读取了 OSGB 数据。
6. 地质图：图例、符号和地层关系。

安全分级：
- 一级：直接观察事实。
- 二级：初步地质解释，必须写依据和置信度。
- 三级：需要验证假设，必须说明无法确定的原因及补充资料。

禁止仅凭照片确定地质时代、精确岩性、断层性质和产状数值。图像质量差时不得强行下结论。
只输出严格 JSON，不要 Markdown，不要思考过程。
""".strip()



def build_prompt(memory_examples, user_focus):
    """把记忆库和本次任务合成提示词。"""
    memory_text = "暂无历史记忆。"
    if memory_examples:
        memory_text = "\n".join(
            [
                f"- 案例{idx + 1}：输入特征={item.get('summary', '无')}；最终结论={item.get('final_conclusion', '无')}"
                for idx, item in enumerate(memory_examples)
            ]
        )

    return f"""
{classify_image_type_prompt()}

历史记忆参考：
{memory_text}

本次用户关注点：{user_focus}

补充要求：
- 自然语言解译必须是适合直接展示给用户的中文自然语言。
- 结构化解译必须是 JSON 对象，按图像类型动态组织字段。
- 如果是 无人机航拍宏观地质影像，结构化解译中建议包含：地形地貌、地层展布、线性构造、沟谷水系、岩性分布。
- 如果是 地面近景地质影像，结构化解译中建议包含：岩性、结构构造、粒度、层理、节理裂隙、风化。
- 如果是 地质剖面图，结构化解译中建议包含：地层接触关系、产状、褶皱、断层、岩体接触、时代关系。
- 如果是 宏微观混合地质影像，结构化解译中建议包含：宏观和微观两个分析块。
- 置信度用 0 到 1 之间的小数表示。
""".strip()



def extract_json_from_text(text):
    """尽量从模型输出中提取 JSON,容错处理多余文本。"""
    return extract_json_object(text)


def prepare_image_bytes(image_bytes, filename):
    """将 TIFF 等格式规范化为 JPEG，供预览和视觉模型稳定使用。"""
    image = Image.open(io.BytesIO(image_bytes)).convert("RGB")
    if image.width > 2048 or image.height > 2048:
        image.thumbnail((2048, 2048))
    output = io.BytesIO()
    image.save(output, format="JPEG", quality=92)
    return output.getvalue(), f"{os.path.splitext(filename)[0]}.jpg"


def build_multiview_prompt(memory_examples, user_focus, photo_records):
    """构造多视角综合判读提示词，要求明确证据来源和待核验项。"""
    memory_text = "暂无历史记忆。"
    if memory_examples:
        memory_text = "\n".join(
            f"- 案例{index + 1}：{item.get('summary', '无')}"
            for index, item in enumerate(memory_examples)
        )
    photo_text = "\n".join(
        f"照片{item['编号']}（{item['文件名']}）：质量={json.dumps(item['图像质量评估'], ensure_ascii=False)}；"
        f"单图解译={json.dumps(item['单图解译'], ensure_ascii=False)}"
        for item in photo_records
    )
    return f"""
你是地质技能大赛 AI 赛道的多视角地质现象综合解译专家。
请基于同一地质现象的多张照片进行交叉核验，只能使用提供资料中可见或明确记录的证据。

历史案例参考：
{memory_text}

本次关注点：{user_focus}

单图解译资料：
{photo_text}

请先以简洁中文给出综合说明，随后输出严格 JSON。JSON 必须包括：
- 图像类型：本案例的主要图像类型
- 图像质量评估：总体质量与限制
- 可见要素：跨照片重复出现的直接可见事实
- 推荐任务
- 不适合任务
- 自然语言解译
- 结构化解译：包含岩性、构造特征、接触关系或地质现象（仅填写证据支持的项目）
- 证据链：数组，每项包含“结论”“证据照片编号”“可见证据”“可靠性”
- 冲突或待核验项：数组，列出照片间不一致或证据不足的判断
- 置信度：0 到 1 的小数
- 不确定性
- 下一步建议

规则：
1. “证据照片编号”只能引用提供的照片编号。
2. 区分直接可见事实和解释结论；无充分证据时写“待核验”，不得臆断断层性质、地层时代或精确岩性。
3. 若有效照片少于 5 张，必须在不确定性中说明多视角证据不足。
4. 不要输出 Markdown。
""".strip()


def make_mock_result(image_quality):
    """离线模式下返回标准化结果，便于比赛现场演示。"""
    return {
        "图像类型": "地面近景地质影像",
        "图像质量评估": image_quality,
        "可见要素": ["岩面", "节理裂隙", "轻微风化面"],
        "推荐研究任务": ["岩性识别", "风化特征识别", "构造特征识别"],
        "不适合研究任务": ["精确层位追踪", "高精度产状测量"],
        "自然语言解译": "这是一个不代表真实地质结论的界面演示结果。当前只记录岩面、裂隙和风化等可见要素，精确岩性及产状均无法确定。",
        "结构化解译": {
            "岩性": "无法确定（模拟界面不提供岩性结论）",
            "结构构造": "可见裂隙，性质待核验",
            "矿物或成分": [],
            "风化程度": "可能存在表面风化，需真实图片确认",
            "地质解译": "需要结合真实照片、比例尺、位置和野外记录。",
        },
        "置信度": 0.0,
        "不确定性": ["缺少比例尺", "缺少拍摄位置和区域地质背景"],
        "下一步建议": ["补充近景与远景各一张", "增加比例尺或地质锤参照", "记录拍摄位置与层位信息"],
        "模型原始输出": "",
    }


# ================= 本地模型调用：所有图片与文字均留在本机 =================
VISION_CONTEXT_WINDOW = 8192
TEXT_CONTEXT_WINDOW = 12288
MAX_MODEL_IMAGE_EDGE = 1024


def stream_local_ollama(
    model_name,
    prompt,
    images=None,
    context_window=TEXT_CONTEXT_WINDOW,
    max_output_tokens=2048,
    keep_alive="5m",
):
    """调用本机 Ollama 的单轮 HTTP 接口，不保留上一轮对话。"""
    payload = {
        "model": model_name,
        "messages": [{"role": "user", "content": prompt, "images": images or []}],
        "stream": True,
        "think": False,
        "keep_alive": keep_alive,
        "options": {
            "temperature": 0.15,
            "num_ctx": context_window,
            "num_predict": max_output_tokens,
        },
    }
    request = urllib.request.Request(
        "http://127.0.0.1:11434/api/chat",
        data=json.dumps(payload).encode("utf-8"),
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    try:
        with urllib.request.urlopen(request, timeout=600) as response:
            for raw_line in response:
                item = json.loads(raw_line.decode("utf-8"))
                text = item.get("message", {}).get("content", "")
                if text:
                    yield text
    except urllib.error.URLError as exc:
        yield f"❌ 无法连接本机 Ollama：{exc}。请确认 Ollama 已启动并已下载模型。"
    except Exception as exc:
        yield f"❌ 本地模型调用异常：{exc}"


def unload_ollama_model(model_name):
    """主动卸载已完成任务的 Ollama 模型，降低 8GB 显存并存风险。"""
    payload = {
        "model": model_name,
        "messages": [],
        "stream": False,
        "keep_alive": 0,
    }
    request = urllib.request.Request(
        "http://127.0.0.1:11434/api/chat",
        data=json.dumps(payload).encode("utf-8"),
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    try:
        with urllib.request.urlopen(request, timeout=60):
            return True
    except Exception:
        # 卸载失败不应覆盖已经完成的地质观察结果。
        return False


def analyze_geology_image(image_path, user_prompt, _api_key, model_name):
    """使用本机视觉模型对单张照片进行初判。"""
    try:
        with open(image_path, "rb") as image_file:
            image = base64.b64encode(image_file.read()).decode("utf-8")
    except FileNotFoundError:
        yield "❌ 找不到图片文件，请检查上传。"
        return
    yield from stream_local_ollama(
        model_name,
        user_prompt,
        [image],
        context_window=VISION_CONTEXT_WINDOW,
        max_output_tokens=900,
        keep_alive="10m",
    )


def render_result(result):#把结构化结果渲染成适合比赛展示的页面。
    """把结构化结果渲染成适合比赛展示的页面。"""
    st.subheader("3. 结构化解译结果")
    if not result:#检查结果
        st.info("请先完成一次解译。")
        return

    if result.get("error"):#检查错误
        st.error(result["error"])
        if result.get("raw_text"):#检查原始输出
            st.text_area("模型原始输出", result["raw_text"], height=180)
        return

    col_a, col_b, col_c = st.columns(3)#创建三个指标展示置信度、图像类型、推荐任务数
    try:
        confidence = min(max(float(result.get("置信度", 0)), 0.0), 1.0)
    except (TypeError, ValueError):
        confidence = 0.0
    col_a.metric("置信度", f"{round(confidence * 100)}%")
    col_b.metric("图像类型", result.get("图像类型", "图像类型不明"))
    col_c.metric("任务数", len(result.get("推荐任务", [])))

    st.text_area("自然语言解译", value=result.get("自然语言解译", ""), height=150)
    evidence_chain = result.get("证据链", [])
    if evidence_chain:
        st.write("证据链（结论必须可回溯至照片编号）")
        st.dataframe(evidence_chain, use_container_width=True)
    verification_items = result.get("冲突或待核验项", [])
    if verification_items:
        st.warning("待核验项：" + "；".join(map(str, verification_items)))
    st.write("可见要素")
    st.write(result.get("可见要素", []))
    st.write("推荐任务")
    st.write(result.get("推荐任务", []))
    st.write("不适合任务")
    st.write(result.get("不适合任务", []))
    st.write("不确定性")
    st.write(result.get("不确定性", []))
    st.write("下一步建议")
    st.write(result.get("下一步建议", []))

    with st.expander("查看底层结构化结果 JSON"):#查看底层结构化结果JSON（给专业人士或调试人员看的）
        st.json(result.get("结构化解译", {}))



def export_docx(image_bytes, result, output_path):#导出Word报告
    """导出 Word 报告；如果缺少依赖，则返回友好提示。"""
    if Document is None:#检查依赖
        return False, "当前环境缺少 python-docx，无法生成 .docx。"

    doc = Document()#创建Word文档
    doc.add_heading("地质AI辅助解译报告", level=1)#添加标题
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")#添加生成时间
    doc.add_paragraph(f"系统版本：{APP_VERSION}")#添加系统版本

    doc.add_heading("一、图像信息", level=2)
    doc.add_paragraph(f"图像类型：{result.get('图像类型', '图像类型不明')}")
    iq = result.get("图像质量评估", {})#获取图像质量评估
    if not isinstance(iq, dict):
        iq = {}#如果图像质量评估不是字典，则设置为空字典
    width = iq.get("图像宽度", iq.get("width", "NA"))#获取图像宽度
    height = iq.get("图像高度", iq.get("height", "NA"))#获取图像高度
    doc.add_paragraph(f"图像宽高：{width} × {height}")

    if image_bytes and Inches is not None:#检查图片
        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
            tmp.write(image_bytes)
            tmp_img_path = tmp.name
        try:#插入图片
            doc.add_picture(tmp_img_path, width=Inches(5.8))
        finally:#清理临时文件
            if os.path.exists(tmp_img_path):#删除临时文件
                os.unlink(tmp_img_path)

    doc.add_heading("二、自然语言结论", level=2)#添加自然语言结论
    doc.add_paragraph(result.get("自然语言解译", ""))

    doc.add_heading("三、结构化解译结果", level=2)#添加结构化解译结果
    structured = result.get("结构化解译", {})#获取结构化解译结果
    if isinstance(structured, dict):#如果结构化解译结果是字典
        for key, value in structured.items():#遍历结构化解译结果
            if isinstance(value, list):
                value_text = ", ".join(map(str, value))#如果结构化解译结果是列表，则将列表转换为字符串
            elif isinstance(value, dict):
                value_text = json.dumps(value, ensure_ascii=False)#如果结构化解译结果是字典，则将字典转换为字符串
            else:
                value_text = str(value)#如果结构化解译结果是其他类型，则将结果转换为字符串
            doc.add_paragraph(f"{key}：{value_text}")
    else:
        doc.add_paragraph(str(structured))#如果结构化解译结果不是字典，则将结果转换为字符串

    doc.add_heading("四、置信度与不确定性", level=2)
    doc.add_paragraph(f"置信度：{result.get('置信度', 0)}")#添加置信度

    uncertainty = result.get("不确定性", [])#获取不确定性
    if isinstance(uncertainty, list):
        uncertainty_text = ", ".join(map(str, uncertainty))#如果不确定性是列表，则将列表转换为字符串
    elif uncertainty is None:
        uncertainty_text = ""#如果不确定性是None，则设置为空字符串
    else:
        uncertainty_text = str(uncertainty)#如果不确定性是其他类型，则将结果转换为字符串    

    suggestions = result.get("下一步建议", [])
    if isinstance(suggestions, list):
        suggestions_text = ", ".join(map(str, suggestions))
    elif suggestions is None:
        suggestions_text = ""
    else:
        suggestions_text = str(suggestions)

    doc.add_paragraph(f"不确定性：{uncertainty_text}")
    doc.add_paragraph(f"建议：{suggestions_text}")

    doc.save(output_path)#保存文档
    return True, output_path#返回成功和路径 (Word报告路径)


def image_bytes_as_jpeg(image_bytes):
    """压缩并缩放图片，避免多视角请求耗尽视觉模型上下文。"""
    image = Image.open(io.BytesIO(image_bytes)).convert("RGB")
    image.thumbnail((MAX_MODEL_IMAGE_EDGE, MAX_MODEL_IMAGE_EDGE))
    output = io.BytesIO()
    image.save(output, format="JPEG", quality=82, optimize=True)
    return output.getvalue()


def build_single_observation_prompt(
    photo_id,
    user_focus,
    observation_point,
    memory_examples=None,
):
    """要求视觉模型输出一张照片的克制、可追溯观测。"""
    memory_text = "暂无历史专家修正案例。"
    if memory_examples:
        memory_text = "\n".join(
            f"- {item.get('观察点编号', '未知点')}：{item.get('summary', '无摘要')}；人工修正={item.get('manual_comment', '无')}"
            for item in memory_examples[-5:]
        )
    return f"""
{classify_image_type_prompt()}

当前观察点编号为“{observation_point}”，照片编号为“{photo_id}”。
用户关注点：{user_focus}
最近专家修正案例（只作表达边界参考，不能复制其地质结论）：
{memory_text}

只分析当前图片。若它只是多视角案例中的一张，图像类型仍按自身内容判断；组合关系由后续文本模型处理。

只输出一个严格 JSON 对象：
{{
  "观察点编号": "{observation_point}",
  "照片编号": "{photo_id}",
  "图像类型": "",
  "直接观察事实": [],
  "岩性外观分析": {{
    "颜色": "",
    "结构": "",
    "粒度": "",
    "矿物外观": ""
  }},
  "构造观察": {{
    "可见现象": "",
    "是否可能存在断裂": "无法确定/可能/未见明显证据",
    "依据": ""
  }},
  "层理信息": "",
  "风化特征": "",
  "可能解释": [{{"结论": "", "依据": "", "置信度": 0.0}}],
  "不可确定内容": [],
  "补充建议": []
}}
内容保持精炼，总长度不超过 700 个汉字。
""".strip()


def build_text_synthesis_prompt(user_focus, single_results, observation_point):
    """将逐图观测整理为纯文本模型可综合的紧凑资料。"""
    observations = []
    for item in single_results:
        initial = item.get("初判", {})
        if not isinstance(initial, dict):
            initial = {"解析失败的原始输出": str(initial)[:800]}
        observations.append({
            "照片编号": item.get("照片编号"),
            "文件名": item.get("文件名"),
            "图像质量统计": item.get("图像质量"),
            "视觉模型观测": initial,
        })
    result_text = json.dumps(observations, ensure_ascii=False, separators=(",", ":"))
    return f"""
你是地质技能大赛野外调查的文本综合分析专家。以下资料是视觉模型对同一地质体各张照片的独立文字观测；你看不到原始图片，只能依据这些记录分析。

观察点编号：{observation_point}
用户关注点：{user_focus}
逐图观测资料：
{result_text}

要求：
1. 区分直接观测事实、地质解释和待验证假设。
2. 每项关键结论必须引用支持它的照片编号；不能声称重新核验了图片。
3. 只有证据充分时才讨论岩性、地层序列、构造关系和演化过程。
4. 不得虚构地层时代、精确岩性、产状数值或断层运动性质。
5. 不同照片记录冲突时必须列入“冲突或待核验项”。
6. 输出一个严格 JSON 对象，不要 Markdown，不要 JSON 前后的说明文字。

JSON 必须包含：
{{
  "观察点编号": "{observation_point}",
  "图像类型": "",
  "图像质量评估": "",
  "可见要素": [],
  "推荐任务": [],
  "不适合任务": [],
  "自然语言解译": "完整但克制的综合地质分析报告",
  "结构化解译": {{
    "岩性分析": "",
    "地层与接触关系": "",
    "构造特征": "",
    "照片间空间与构造联系": "",
    "可能的形成或演化过程": ""
  }},
  "证据链": [{{"结论": "", "观察点编号": ["{observation_point}"], "证据照片编号": [], "观测依据": "", "可靠性": ""}}],
  "冲突或待核验项": [],
  "置信度": 0.0,
  "不确定性": [],
  "下一步建议": []
}}
""".strip()


def validate_single_observation(result, observation_point, photo_id):
    """校验单图观察的绑定编号和核心字段。"""
    if not isinstance(result, dict):
        return ["单图结果不是 JSON 对象"]
    errors = []
    if result.get("观察点编号") != observation_point:
        errors.append("观察点编号不匹配")
    if result.get("照片编号") != photo_id:
        errors.append("照片编号不匹配")
    for key, expected_type in {
        "图像类型": str,
        "直接观察事实": list,
        "岩性外观分析": dict,
        "构造观察": dict,
        "可能解释": list,
        "不可确定内容": list,
        "补充建议": list,
    }.items():
        if not isinstance(result.get(key), expected_type):
            errors.append(f"字段缺失或类型错误：{key}")
    return errors


def validate_synthesis_result(result, observation_point, valid_photo_ids):
    """校验多视角综合结果及其照片证据引用。"""
    if not isinstance(result, dict):
        return ["综合结果不是 JSON 对象"]
    errors = []
    if result.get("观察点编号") != observation_point:
        errors.append("综合结果的观察点编号不匹配")
    for key, expected_type in {
        "自然语言解译": str,
        "结构化解译": dict,
        "证据链": list,
        "不确定性": list,
        "下一步建议": list,
    }.items():
        if not isinstance(result.get(key), expected_type):
            errors.append(f"字段缺失或类型错误：{key}")
    try:
        confidence = float(result.get("置信度"))
        if not 0 <= confidence <= 1:
            raise ValueError
    except (TypeError, ValueError):
        errors.append("置信度必须是 0 到 1 之间的小数")
    for index, evidence in enumerate(result.get("证据链", []) if isinstance(result.get("证据链"), list) else []):
        if not isinstance(evidence, dict):
            errors.append(f"证据链第{index + 1}项不是对象")
            continue
        point_refs = evidence.get("观察点编号", [])
        if isinstance(point_refs, str):
            point_refs = [point_refs]
        if set(point_refs or []) != {observation_point}:
            errors.append(f"证据链第{index + 1}项观察点引用错误")
        photo_refs = evidence.get("证据照片编号", [])
        if isinstance(photo_refs, str):
            photo_refs = [photo_refs]
        unknown = set(photo_refs or []) - set(valid_photo_ids)
        if unknown:
            errors.append(f"证据链第{index + 1}项引用未知照片：{sorted(unknown)}")
    return errors


def synthesize_observations(user_prompt, model_name):
    """使用独立纯文本模型综合逐图观测，不再传递任何图片。"""
    yield from stream_local_ollama(
        model_name,
        user_prompt,
        images=None,
        context_window=TEXT_CONTEXT_WINDOW,
        max_output_tokens=2400,
        keep_alive=0,
    )


def export_project_docx(project, database, output_path):
    """从项目 JSON 和数据库生成 5000 字以内的竞赛调查报告。"""
    if Document is None:
        return False, "当前环境缺少 python-docx，无法生成 .docx。"

    doc = Document()
    doc.add_heading("基于离线多模态AI的地质调查报告", level=1)
    doc.add_paragraph(f"项目名称：{project.get('项目名称', '')}")
    doc.add_paragraph(f"调查路线：{project.get('路线', '')}")
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"系统版本：{APP_VERSION}")

    report_parts = []
    model_intro = (
        "系统完全离线运行，使用 Ollama 部署 qwen3-vl:8b 作为野外地质观察助手，"
        "逐图提取可见事实；使用 qwen3:8b 对结构化观察记录进行多点综合分析。"
        "程序负责项目数据组织、证据校验和剖面绘制，模型负责观察、解释与推理。"
    )
    doc.add_heading("一、AI模型介绍", level=2)
    doc.add_paragraph(model_intro)
    report_parts.append(model_intro)

    process_text = (
        "原始照片按项目、观察点和稳定照片编号归档；视觉模型逐图输出 JSON，随后由文本模型"
        "综合同一观察点和多个观察点之间的关系。所有结论通过观察点编号与照片编号回溯，"
        "无法从资料确定的时代、精确岩性、产状和断层性质保留为不确定项。"
    )
    doc.add_heading("二、数据处理流程与质量控制", level=2)
    doc.add_paragraph(process_text)
    report_parts.append(process_text)

    doc.add_heading("三、观察点空间拓扑与地质现象识别", level=2)
    route_records = ordered_spatial_records(database)
    doc.add_paragraph("以下路线顺序、点间距离、方位、高程和高程差均由地质人员人工输入；程序仅负责按人工顺序整理，未进行自动排序或空间计算。")
    for route_record in route_records:
        point_id = route_record["观察点编号"]
        record = database.get(point_id, {})
        doc.add_heading(point_id, level=3)
        spatial_text = json.dumps(route_record, ensure_ascii=False)
        doc.add_paragraph(f"空间信息：{spatial_text}")
        result = record.get("视觉观察", {})
        conclusion = ""
        if isinstance(result, dict):
            conclusion = str(result.get("自然语言解译") or result.get("结构化解译") or "")
        conclusion = conclusion[:1200]
        doc.add_paragraph(f"AI观察与解释：{conclusion or '尚未完成AI分析。'}")
        correction = str(record.get("人工修正", ""))
        if correction:
            doc.add_paragraph(f"人工修正：{correction}")
        report_parts.extend([spatial_text, conclusion, correction])

    doc.add_heading("四、区域综合地质解释", level=2)
    regional = project.get("区域综合分析", {})
    regional_text = (
        json.dumps(regional, ensure_ascii=False, indent=2)[:1800]
        if regional
        else "尚未完成区域综合分析。"
    )
    doc.add_paragraph(regional_text)
    report_parts.append(regional_text)

    profile_result = project.get("剖面成果", {})
    profile_png = (
        profile_result.get("输出文件", {}).get("PNG")
        if isinstance(profile_result, dict)
        else None
    )
    if profile_png and Path(profile_png).exists() and Inches is not None:
        doc.add_heading("五、AI辅助地质剖面", level=2)
        doc.add_picture(profile_png, width=Inches(6.2))
        doc.add_paragraph("剖面由结构化参数经人工复核后使用 Matplotlib 程序化绘制，不使用生图模型。")

    innovation = (
        "主要创新点包括：小模型本地离线部署；观察点—照片—事实—结论四级证据链；"
        "多视角与多观察点结构化融合；地质不确定性约束；AI参数与程序化地质剖面绘制分工。"
    )
    doc.add_heading("六、主要创新点", level=2)
    doc.add_paragraph(innovation)
    report_parts.append(innovation)

    approximate_count = len("".join(report_parts))
    if approximate_count > 5000:
        return False, f"报告正文约 {approximate_count} 字，超过 5000 字限制，请精简后导出。"
    doc.add_paragraph(f"正文统计：约 {approximate_count} 字。")
    doc.save(output_path)
    return True, output_path


def ordered_spatial_records(database_data):
    """按人工填写的观察点顺序整理路线，并兼容当前数据库空间字段。"""
    records = []
    all_topology = database_data.get("observation_topology", {})
    for point_id, record in database_data.items():
        if point_id == "observation_topology" or not isinstance(record, dict):
            continue
        spatial = record.get("空间信息", {})
        topology = all_topology.get(point_id, record.get("空间拓扑", {}))
        records.append({
            "观察点编号": point_id,
            "观察点顺序": spatial.get("观察点顺序", topology.get("观察点顺序")),
            "下一观察点编号": spatial.get(
                "下一观察点编号",
                topology.get("下一观察点编号", ""),
            ),
            "点间距离（m）": spatial.get(
                "DasViewer测量距离",
                spatial.get("距离下一观察点", topology.get("点间距离（m）")),
            ),
            "方位": spatial.get(
                "DasViewer测量方位",
                spatial.get("方位", topology.get("方位（°）", "")),
            ),
            "高程（m）": spatial.get(
                "模型坐标Z",
                spatial.get("高程", topology.get("高程")),
            ),
            "高程差（m）": spatial.get(
                "高程差",
                topology.get("高程差"),
            ),
            "空间关系描述": spatial.get(
                "空间关系描述",
                spatial.get("空间关系说明", topology.get("空间关系描述", "")),
            ),
            "数据来源": "地质人员人工输入或根据模型坐标自动计算",
        })
    return sorted(records, key=lambda item: (
        item["观察点顺序"] is None,
        item["观察点顺序"] if item["观察点顺序"] is not None else float("inf"),
        item["观察点编号"],
    ))


def parse_optional_float(value, field_name):
    """把界面中的可选数值转换为浮点数，空值保留为 None。"""
    text = str(value or "").strip()
    if not text:
        return None
    try:
        return float(text)
    except ValueError as exc:
        raise ValueError(f"{field_name} 必须是数字或留空。") from exc


# ================= 主界面 =================
st.title("🪨 地质AI辅助解译与成图系统")
st.caption("V4.0：项目级证据管理、多观察点综合分析、程序化地质剖面和报告导出。")
st.markdown("---")

# 侧边栏设置
st.sidebar.title("⚙️ 系统设置")
local_base_url = "http://127.0.0.1:11434"
vision_model_choice = st.sidebar.selectbox("单图视觉模型", ["qwen3-vl:8b"])
text_model_choice = st.sidebar.selectbox("文本综合模型", ["qwen3:8b"])
user_focus = st.sidebar.text_area("本次关注点", "请在此输入您本次的解译关注点")
selected_page = st.sidebar.radio(
    "功能页面",
    ["项目数据管理", "照片AI识别", "区域综合分析", "AI剖面图", "报告导出"],
)
st.sidebar.info(
    f"当前系统版本：{APP_VERSION}\n"
    "处理流程：视觉模型逐图观测 → 文本模型综合报告\n"
    "本地 Ollama 模式：不使用外网"
)

project_summaries = PROJECT_MANAGER.list_projects()
project_keys = [item["项目标识"] for item in project_summaries]
project_labels = {
    item["项目标识"]: f"{item['项目名称']}（{item['观察点数量']}个观察点）"
    for item in project_summaries
}
current_project_key = None
current_project = None
current_point_id = None

if project_keys:
    previous_key = st.session_state.get("current_project_key")
    default_index = project_keys.index(previous_key) if previous_key in project_keys else 0
    current_project_key = st.sidebar.selectbox(
        "当前项目",
        project_keys,
        index=default_index,
        format_func=lambda key: project_labels.get(key, key),
    )
    st.session_state.current_project_key = current_project_key
    try:
        current_project = PROJECT_MANAGER.load_project(current_project_key)
    except ProjectManagerError as exc:
        st.sidebar.error(str(exc))
    if current_project:
        point_ids = [item.get("编号") for item in current_project.get("观察点", []) if item.get("编号")]
        if point_ids:
            current_point_id = st.sidebar.selectbox("当前观察点", point_ids)


if selected_page == "项目数据管理":
    st.header("1. 项目数据管理")
    with st.form("create_project_form"):
        new_project_name = st.text_input("新项目名称")
        new_project_route = st.text_input("调查路线")
        create_submitted = st.form_submit_button("创建项目", type="primary")
    if create_submitted:
        try:
            new_key, _ = PROJECT_MANAGER.create_project(new_project_name, new_project_route)
            st.session_state.current_project_key = new_key
            st.success(f"项目已创建：{new_project_name}")
            st.rerun()
        except ProjectManagerError as exc:
            st.error(str(exc))

    if current_project_key and current_project:
        col_a, col_b, col_c = st.columns(3)
        col_a.metric("项目", current_project.get("项目名称", ""))
        col_b.metric("观察点", len(current_project.get("观察点", [])))
        col_c.metric("数据版本", current_project.get("数据版本", "4.0"))

        try:
            topology_database = GeologyDatabase(PROJECT_MANAGER.project_dir(current_project_key) / "geology_database.json")
            topology_database.rebuild_observation_topology()
            topology_records = list(topology_database.load().get("observation_topology", {}).values())
            st.subheader("人工调查路线与空间拓扑")
            st.caption("空间事实全部由地质人员输入；程序仅按‘观察点顺序’展示和传递，不自动修改、计算或猜测。")
            st.dataframe(topology_records, use_container_width=True, hide_index=True)
        except (ProjectManagerError, GeologyDatabaseError) as exc:
            st.warning(f"空间拓扑暂时无法读取：{exc}")

        if st.button("新增观察点", use_container_width=True):
            try:
                observation = PROJECT_MANAGER.add_observation(current_project_key)
                database = GeologyDatabase(PROJECT_MANAGER.project_dir(current_project_key) / "geology_database.json")
                database.ensure_observation(observation["编号"], observation["空间信息"])
                st.success(f"已新增观察点 {observation['编号']}")
                st.rerun()
            except (ProjectManagerError, GeologyDatabaseError) as exc:
                st.error(str(exc))

        if current_point_id:
            observation = PROJECT_MANAGER.get_observation(current_project_key, current_point_id)
            spatial = observation.get("空间信息", {})
            st.subheader(f"观察点 {current_point_id}")
            with st.form("observation_form"):
                point_order = st.text_input("观察点顺序", value=str(spatial.get("观察点顺序") or ""), help="由地质人员人工填写；程序仅按该顺序管理相邻关系。")
                next_point_id = st.text_input("下一观察点编号", value=str(spatial.get("下一观察点编号") or ""), help="由地质人员指定相邻观察点；距离、方位和高差由 XYZ 坐标自动计算。")
                coordinate = st.text_input("坐标", value=str(spatial.get("坐标") or ""))
                slope = st.text_input("坡度（°）", value=str(spatial.get("坡度") or ""))
                aspect = st.text_input("坡向", value=str(spatial.get("坡向") or ""))
                distance_next = st.text_input(
                    "距离下一观察点（m）",
                    value=str(spatial.get("距离下一观察点") or ""),
                    help="可填写 DasViewer 人工测量结果；留空表示暂不提供该距离，不从截图或模型自动猜测。",
                )
                st.caption("DasViewer 人工空间测量")
                model_x = st.text_input("模型坐标 X", value=str(spatial.get("模型坐标X") or ""))
                model_y = st.text_input("模型坐标 Y", value=str(spatial.get("模型坐标Y") or ""))
                model_z = st.text_input("模型坐标 Z / 地形高程（m）", value=str(spatial.get("模型坐标Z") or ""))
                model_unit = st.text_input("模型坐标单位", value=str(spatial.get("模型坐标单位") or "m"))
                measured_distance = st.text_input(
                    "DasViewer 测量距离（m）",
                    value=str(spatial.get("DasViewer测量距离") or ""),
                    help="如果填写，将作为人工确认的观察点间距离。",
                )
                measured_azimuth = st.text_input(
                    "DasViewer 测量方位",
                    value=str(spatial.get("DasViewer测量方位") or ""),
                    help="填写模型中观察点指向下一观察点的方位，例如 N、NE、135°。",
                )
                spatial_relation = st.text_area(
                    "空间关系描述",
                    value=str(spatial.get("空间关系描述", spatial.get("空间关系说明", "")) or ""),
                    help="填写 DasViewer 中观察点之间的地形、空间位置、连续性或遮挡关系。",
                    height=90,
                )
                auxiliary_text = st.text_area(
                    "三维辅助信息 JSON（只填写截图说明和人工测量参数，不直接读取 OSGB）",
                    value=json.dumps(observation.get("三维辅助信息", {}), ensure_ascii=False, indent=2),
                    height=150,
                )
                correction = st.text_area("人工修正", value=observation.get("人工修正", ""))
                review_status = st.selectbox(
                    "审核状态",
                    ["未审核", "已确认", "已修正"],
                    index=["未审核", "已确认", "已修正"].index(
                        observation.get("审核状态", "未审核")
                        if observation.get("审核状态", "未审核") in ["未审核", "已确认", "已修正"]
                        else "未审核"
                    ),
                )
                observation_submitted = st.form_submit_button("保存观察点")
            if observation_submitted:
                try:
                    spatial_info = {
                        "观察点顺序": parse_optional_float(point_order, "观察点顺序"),
                        "下一观察点编号": next_point_id.strip().upper(),
                        "坐标": coordinate.strip(),
                        "坡度": parse_optional_float(slope, "坡度"),
                        "坡向": aspect.strip(),
                        "距离下一观察点": parse_optional_float(distance_next, "距离下一观察点"),
                        "距离来源": "人工输入" if str(distance_next or "").strip() else "",
                        "距离计算方式": "人工输入或 DasViewer 测量" if str(distance_next or "").strip() else "",
                        "模型坐标X": parse_optional_float(model_x, "模型坐标 X"),
                        "模型坐标Y": parse_optional_float(model_y, "模型坐标 Y"),
                        "模型坐标Z": parse_optional_float(model_z, "模型坐标 Z"),
                        "模型坐标单位": model_unit.strip(),
                        "DasViewer测量距离": parse_optional_float(measured_distance, "DasViewer 测量距离"),
                        "DasViewer测量方位": measured_azimuth.strip(),
                        "空间关系描述": spatial_relation.strip(),
                        "空间关系说明": spatial_relation.strip(),
                    }
                    if spatial_info["DasViewer测量距离"] is not None:
                        spatial_info["距离来源"] = "人工输入"
                        spatial_info["距离计算方式"] = "DasViewer 人工测量"
                    auxiliary = json.loads(auxiliary_text or "{}")
                    if not isinstance(auxiliary, dict):
                        raise ValueError("三维辅助信息必须是 JSON 对象。")
                    PROJECT_MANAGER.update_observation(
                        current_project_key,
                        current_point_id,
                        spatial_info=spatial_info,
                        auxiliary_3d=auxiliary,
                        manual_correction=correction,
                        review_status=review_status,
                    )
                    database = GeologyDatabase(PROJECT_MANAGER.project_dir(current_project_key) / "geology_database.json")
                    database.ensure_observation(current_point_id, spatial_info)
                    database.rebuild_observation_topology()
                    database.update_manual_correction(current_point_id, correction, review_status)
                    st.success("观察点已保存。")
                    st.rerun()
                except (ValueError, json.JSONDecodeError, ProjectManagerError, GeologyDatabaseError) as exc:
                    st.error(str(exc))

        with st.expander("查看 project.json"):
            st.json(current_project)
        try:
            database = GeologyDatabase(PROJECT_MANAGER.project_dir(current_project_key) / "geology_database.json")
            with st.expander("查看 geology_database.json"):
                st.json(database.load())
        except (ProjectManagerError, GeologyDatabaseError) as exc:
            st.error(str(exc))
    else:
        st.info("请先创建一个项目。")
    st.stop()


if selected_page == "区域综合分析":
    st.header("3. 区域综合分析")
    if not current_project_key or not current_project:
        st.warning("请先创建并选择项目。")
        st.stop()
    try:
        database = GeologyDatabase(PROJECT_MANAGER.project_dir(current_project_key) / "geology_database.json")
        payload = database.build_regional_payload()
        point_ids_for_analysis = list(payload.keys())
        st.caption(
            f"当前项目共有 {len(payload)} 个观察点，区域综合分析将纳入全部观察点。"
            "区域分析至少需要 2 个观察点。"
        )
        st.info("本次将分析观察点：" + "、".join(point_ids_for_analysis))
        st.dataframe(
            [value.get("路线节点", {}) for value in payload.values()],
            use_container_width=True,
            hide_index=True,
        )
        if st.button("开始区域综合分析", type="primary", disabled=len(payload) < 2):
            with st.spinner("qwen3:8b 正在分析多个观察点之间的关系..."):
                regional_result = analyze_region(
                    current_project,
                    payload,
                    model_name=text_model_choice,
                    base_url=local_base_url,
                )
                PROJECT_MANAGER.set_regional_analysis(current_project_key, regional_result)
                current_project = PROJECT_MANAGER.load_project(current_project_key)
                st.success("区域综合分析已保存。")
        regional_result = current_project.get("区域综合分析", {})
        if regional_result:
            st.metric("区域置信度", f"{float(regional_result.get('置信度', 0)) * 100:.0f}%")
            st.write(regional_result.get("区域地质概况", ""))
            st.subheader("观察点覆盖情况")
            st.dataframe(regional_result.get("观察点覆盖情况", []), use_container_width=True)
            st.subheader("证据链")
            st.dataframe(regional_result.get("证据链", []), use_container_width=True)
            with st.expander("完整区域分析 JSON"): 
                st.json(regional_result)
    except (ProjectManagerError, GeologyDatabaseError, RegionalAnalysisError) as exc:
        st.error(str(exc))
    st.stop()


if selected_page == "AI剖面图":
    st.header("4. AI辅助地质剖面图")
    if not current_project_key or not current_project:
        st.warning("请先创建并选择项目。")
        st.stop()
    try:
        database = GeologyDatabase(PROJECT_MANAGER.project_dir(current_project_key) / "geology_database.json")
        payload = database.build_regional_payload()
        regional_result = current_project.get("区域综合分析", {})
        st.caption("剖面参数只能使用人工空间拓扑中的顺序、点间距离、方位、高程和高程差；程序与模型均不得补造空间数值。")
        st.dataframe(
            [value.get("路线节点", {}) for value in payload.values()],
            use_container_width=True,
            hide_index=True,
        )
        profile_direction = st.text_input("剖面方向", value="NW-SE")
        profile_state_key = f"profile_json_{current_project_key}"
        if profile_state_key not in st.session_state:
            existing_parameters = current_project.get("剖面成果", {}).get("参数", {})
            st.session_state[profile_state_key] = json.dumps(existing_parameters, ensure_ascii=False, indent=2) if existing_parameters else "{}"

        if st.button("AI生成剖面参数草案", type="primary"):
            with st.spinner("正在依据结构化观察记录整理剖面参数..."):
                draft = generate_profile_parameters(
                    payload,
                    regional_result,
                    profile_direction,
                    model_name=text_model_choice,
                    base_url=local_base_url,
                )
                st.session_state[profile_state_key] = json.dumps(draft, ensure_ascii=False, indent=2)
                if not draft.get("数据充分性"):
                    st.warning("数据不足：" + "；".join(map(str, draft.get("无法生成原因", []))))

        profile_text = st.text_area(
            "剖面参数 JSON（绘图前必须人工复核）",
            value=st.session_state[profile_state_key],
            height=420,
        )
        if st.button("校验并绘制 PNG/SVG", use_container_width=True):
            profile_data = json.loads(profile_text)
            errors = validate_profile_data(profile_data)
            if errors:
                st.error("；".join(errors))
            else:
                output_files = draw_geological_profile(
                    profile_data,
                    PROJECT_MANAGER.project_dir(current_project_key) / "exports",
                    basename=f"{current_project_key}_geology_profile",
                )
                result = {"参数": profile_data, "输出文件": output_files, "人工确认时间": datetime.now().isoformat()}
                PROJECT_MANAGER.set_profile_result(current_project_key, result)
                st.success("剖面图已生成并写入项目。")
                st.image(output_files["PNG"], use_container_width=True)

        saved_profile = current_project.get("剖面成果", {})
        saved_files = saved_profile.get("输出文件", {}) if isinstance(saved_profile, dict) else {}
        if saved_files.get("PNG") and Path(saved_files["PNG"]).exists():
            st.image(saved_files["PNG"], caption="当前项目剖面成果", use_container_width=True)
            for file_type in ("PNG", "SVG"):
                file_path = saved_files.get(file_type)
                if file_path and Path(file_path).exists():
                    st.download_button(
                        f"下载 {file_type}",
                        data=Path(file_path).read_bytes(),
                        file_name=Path(file_path).name,
                        mime="image/png" if file_type == "PNG" else "image/svg+xml",
                    )
    except (json.JSONDecodeError, ProjectManagerError, GeologyDatabaseError, ProfileGenerationError, RegionalAnalysisError) as exc:
        st.error(str(exc))
    st.stop()


if selected_page == "报告导出":
    st.header("5. 报告导出")
    if not current_project_key or not current_project:
        st.warning("请先创建并选择项目。")
        st.stop()
    try:
        database = GeologyDatabase(PROJECT_MANAGER.project_dir(current_project_key) / "geology_database.json")
        database_data = database.load()
        if st.button("生成项目级 AI 地质调查报告", type="primary", use_container_width=True):
            output_path = PROJECT_MANAGER.project_dir(current_project_key) / "exports" / f"{current_project_key}_AI地质调查报告.docx"
            ok, result = export_project_docx(current_project, database_data, output_path)
            if ok:
                st.download_button(
                    "下载 Word 报告",
                    data=Path(result).read_bytes(),
                    file_name=Path(result).name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
                st.success("项目级报告已生成。")
            else:
                st.warning(result)
        st.download_button(
            "下载核心数据库 JSON",
            data=json.dumps(database_data, ensure_ascii=False, indent=2).encode("utf-8"),
            file_name=f"{current_project_key}_geology_database.json",
            mime="application/json",
            use_container_width=True,
        )
    except (ProjectManagerError, GeologyDatabaseError, OSError) as exc:
        st.error(str(exc))
    st.stop()


if not current_project_key or not current_project or not current_point_id:
    st.warning("照片 AI 识别必须绑定项目和观察点。请先在“项目数据管理”页面创建项目与观察点。")
    st.stop()

analysis_context = f"{current_project_key}:{current_point_id}"
if st.session_state.get("analysis_context") != analysis_context:
    st.session_state.analysis_context = analysis_context
    st.session_state.raw_stream_text = ""
    st.session_state.structured_result = {}
    current_observation = PROJECT_MANAGER.get_observation(current_project_key, current_point_id)
    st.session_state.manual_comment = current_observation.get("人工修正", "")
    st.session_state.manual_lithology = ""

col1, col2 = st.columns([1, 1.25])

with col1:
    st.subheader(f"📥 2. 照片 AI 识别 · {current_point_id}")
    st.caption("请将同一地质现象的照片作为一个案例上传，建议至少提供 5 个不同视角。")
    uploader_key = f"photo_uploader_{current_project_key}_{current_point_id}"
    uploaded_files = st.file_uploader(
        "上传同一地质现象的多视角照片",
        type=["jpg", "png", "jpeg", "tif", "tiff"],
        accept_multiple_files=True,
        key=uploader_key,
    )
    image_items = []
    image_quality = {}

    if uploaded_files:
        for index, uploaded in enumerate(uploaded_files, start=1):
            raw_bytes = uploaded.getvalue()
            try:
                quality = get_image_quality(raw_bytes)
                image_items.append({
                    "id": f"{current_point_id}-待上传-{index:02d}",
                    "name": uploaded.name,
                    "bytes": raw_bytes,
                    "quality": quality,
                    "record": None,
                })
            except (ProjectManagerError, OSError, ValueError) as exc:
                st.warning(f"无法读取 {uploaded.name}：{exc}")

        st.info(f"已载入 {len(image_items)} 张照片。")
        if len(image_items) < 5:
            st.warning("当前照片数量少于 5 张，多视角证据不足，建议补充同一现象的其他角度。")
        preview_columns = st.columns(min(len(image_items), 5))
        for index, item in enumerate(image_items):
            with preview_columns[index % len(preview_columns)]:
                st.image(item["bytes"], caption=item["id"], use_container_width=True)
                with st.expander("质量", expanded=False):
                    st.json(item["quality"])
        image_quality = image_items[0]["quality"] if image_items else {}

    st.subheader("✏️ 4. 人工修正")
    manual_lithology = st.text_input("修正岩性", value=st.session_state.get("manual_lithology", ""))#修正岩性
    manual_comment = st.text_area("修正说明", value=st.session_state.get("manual_comment", ""), height=90)#修正说明
    if st.button("保存人工修正", use_container_width=True):#保存人工修正
        st.session_state.manual_lithology = manual_lithology
        st.session_state.manual_comment = manual_comment
        correction_text = "；".join(
            item for item in [f"修正岩性：{manual_lithology}" if manual_lithology else "", manual_comment]
            if item
        )
        try:
            PROJECT_MANAGER.update_observation(
                current_project_key,
                current_point_id,
                manual_correction=correction_text,
                review_status="已修正" if correction_text else "未审核",
            )
            database = GeologyDatabase(PROJECT_MANAGER.project_dir(current_project_key) / "geology_database.json")
            database.update_manual_correction(
                current_point_id,
                correction_text,
                "已修正" if correction_text else "未审核",
            )
            st.success("人工修正已立即写入项目和地质数据库。")
        except (ProjectManagerError, GeologyDatabaseError) as exc:
            st.error(str(exc))

with col2:
    st.subheader("🤖 2. AI 智能解译")

    if "raw_stream_text" not in st.session_state:#初始化session state(会话状态)
        st.session_state.raw_stream_text = ""
    if "structured_result" not in st.session_state:
        st.session_state.structured_result = {}
    if "tmp_path" not in st.session_state:
        st.session_state.tmp_path = None

    if image_items:
        if st.button("🚀 开始双模型地质解译", type="primary", use_container_width=True):
            st.session_state.raw_stream_text = ""
            st.session_state.structured_result = {}
            try:
                saved_records = PROJECT_MANAGER.add_photos(
                    current_project_key,
                    current_point_id,
                    [(item["name"], item["bytes"]) for item in image_items],
                )
                for item, record in zip(image_items, saved_records):
                    item["id"] = record["照片编号"]
                    item["record"] = record
                saved_observation = PROJECT_MANAGER.get_observation(
                    current_project_key,
                    current_point_id,
                )
                database = GeologyDatabase(
                    PROJECT_MANAGER.project_dir(current_project_key) / "geology_database.json"
                )
                database.ensure_observation(
                    current_point_id,
                    saved_observation.get("空间信息", {}),
                )
                database.update_photo_evidence(
                    current_point_id,
                    saved_observation.get("照片元数据", []),
                )
            except (ProjectManagerError, GeologyDatabaseError) as exc:
                st.error(str(exc))
                st.stop()

            progress = st.progress(0, text="准备逐图提取地质观测...")
            single_results = []
            project_memory_file = PROJECT_MANAGER.project_dir(current_project_key) / "geo_memory.jsonl"
            memory_examples = load_memory_examples(limit=5, memory_file=project_memory_file)
            for index, item in enumerate(image_items, start=1):
                progress.progress(
                    (index - 1) / (len(image_items) + 1),
                    text=f"视觉模型正在分析 {item['id']}（{index}/{len(image_items)}）",
                )
                with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as tmp_file:
                    tmp_file.write(image_bytes_as_jpeg(item["bytes"]))
                    single_path = tmp_file.name
                try:
                    single_prompt = build_single_observation_prompt(
                        item["id"],
                        user_focus,
                        current_point_id,
                        memory_examples,
                    )
                    single_text = "".join(
                        analyze_geology_image(
                            single_path,
                            single_prompt,
                            local_base_url,
                            vision_model_choice,
                        )
                    )
                    parsed_single = extract_json_from_text(single_text)
                    single_errors = validate_single_observation(
                        parsed_single,
                        current_point_id,
                        item["id"],
                    )
                    single_results.append({
                        "照片编号": item["id"],
                        "文件名": item["name"],
                        "图像质量": item["quality"],
                        "初判": parsed_single
                        if parsed_single is not None and not single_errors
                        else {
                            "校验错误": single_errors,
                            "解析失败的原始输出": single_text[:1200],
                        },
                    })
                finally:
                    if os.path.exists(single_path):
                        os.unlink(single_path)

            unload_ollama_model(vision_model_choice)
            progress.progress(
                len(image_items) / (len(image_items) + 1),
                text=f"文本模型 {text_model_choice} 正在综合逐图观测...",
            )
            combined_prompt = build_text_synthesis_prompt(
                user_focus,
                single_results,
                current_point_id,
            )
            full_text = st.write_stream(synthesize_observations(combined_prompt, text_model_choice))
            st.session_state.raw_stream_text = full_text or ""
            parsed = extract_json_from_text(st.session_state.raw_stream_text)
            synthesis_errors = validate_synthesis_result(
                parsed,
                current_point_id,
                {item["id"] for item in image_items},
            )
            st.session_state.structured_result = (
                parsed
                if parsed is not None and not synthesis_errors
                else {
                    "error": "文本模型输出未通过 JSON 与证据引用校验："
                    + "；".join(synthesis_errors),
                    "raw_text": st.session_state.raw_stream_text,
                }
            )
            st.session_state.structured_result["单图初判"] = single_results
            progress.progress(1.0, text="双模型地质解译完成")

            if st.session_state.structured_result and not st.session_state.structured_result.get("error"):
                st.session_state.structured_result["观察点编号"] = current_point_id
                st.session_state.structured_result["manual_lithology"] = st.session_state.get("manual_lithology", "")
                st.session_state.structured_result["manual_comment"] = st.session_state.get("manual_comment", "")
                st.session_state.structured_result["image_quality_assessment"] = image_quality or {}
                correction_text = "；".join(
                    item
                    for item in [
                        f"修正岩性：{st.session_state.get('manual_lithology', '')}"
                        if st.session_state.get("manual_lithology", "")
                        else "",
                        st.session_state.get("manual_comment", ""),
                    ]
                    if item
                )
                observation = PROJECT_MANAGER.update_observation(
                    current_project_key,
                    current_point_id,
                    ai_result=st.session_state.structured_result,
                    manual_correction=correction_text,
                )
                database = GeologyDatabase(PROJECT_MANAGER.project_dir(current_project_key) / "geology_database.json")
                database.ensure_observation(current_point_id, observation.get("空间信息", {}))
                database.update_photo_evidence(
                    current_point_id,
                    observation.get("照片元数据", []),
                )
                database.update_visual_observation(
                    current_point_id,
                    st.session_state.structured_result,
                )
                database.update_manual_correction(
                    current_point_id,
                    correction_text,
                    observation.get("审核状态", "未审核"),
                )
                save_memory_record({
                    "timestamp": datetime.now().isoformat(),
                    "观察点编号": current_point_id,
                    "summary": (st.session_state.structured_result.get("自然语言解译") or "")[:240],
                    "final_conclusion": json.dumps(st.session_state.structured_result, ensure_ascii=False),
                    "manual_lithology": st.session_state.get("manual_lithology", ""),
                    "manual_comment": st.session_state.get("manual_comment", ""),
                }, memory_file=project_memory_file)
                st.success(f"{current_point_id} 多视角解译完成，结果已写入项目、数据库和记忆库。")

        if st.session_state.raw_stream_text:
            st.subheader("3. 多视角综合解译原文")
            st.text_area("实时生成内容", st.session_state.raw_stream_text, height=180)

        if st.session_state.structured_result:
            render_result(st.session_state.structured_result)
            if st.session_state.structured_result.get("证据链"):
                st.subheader("证据链核验")
                st.json(st.session_state.structured_result["证据链"])
            if st.session_state.structured_result.get("单图初判"):
                with st.expander("查看逐图初判结果"):
                    st.json(st.session_state.structured_result["单图初判"])

# ================= 导出 Word 报告 =================
st.markdown("---")
if image_items and st.session_state.structured_result and not st.session_state.structured_result.get("error"): 
    st.subheader("5. 报告导出")
    if st.button("📥 导出标准化地质报告 (Word)", use_container_width=True):
        with st.spinner("正在生成 Word 报告..."):
            tmp_docx = os.path.join(
                tempfile.gettempdir(),
                f"geology_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            )
            ok, result = export_docx(image_items[0]["bytes"], st.session_state.structured_result, tmp_docx)
            if ok:
                with open(result, "rb") as f:
                    st.download_button(
                        "点击下载 Word 报告",
                        data=f.read(),
                        file_name=os.path.basename(result),
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )
                st.toast("Word 报告已生成，可以下载。", icon="📄")
            else:
                st.warning(result)
